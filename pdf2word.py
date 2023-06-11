from pdfminer.high_level import extract_text, extract_pages
from pdfminer.layout import LTImage
from docx import Document
from docx.shared import Inches

# Open PDF file and extract text to string
with open('example.pdf', 'rb') as pdf_file:
    pdf_text = extract_text(pdf_file)
    
# Create new Word document
doc = Document()

# Add text to Word document
doc.add_paragraph(pdf_text)

# Extract images from PDF and add to Word document
for page_layout in extract_pages(pdf_file):
    for element in page_layout:
        if isinstance(element, LTImage):
            # Get image data
            image_data = element.stream.get_rawdata()
            
            # Add image to Word document
            doc.add_picture(image_data, width=Inches(element.width/72), height=Inches(element.height/72))

# Save Word document
doc.save('example.docx')
